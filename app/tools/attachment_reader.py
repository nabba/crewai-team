"""
attachment_reader.py — Extract text content from uploaded Signal attachments.

Supports: PDF, DOCX, XLSX, PNG, JPG (OCR via pytesseract if available,
          otherwise basic image metadata).

Attachments arrive from signal-cli at /app/attachments/<id> and are
mounted read-only into the container.
"""

import logging
import pathlib
from crewai.tools import tool
from app.audit import log_tool_blocked

logger = logging.getLogger(__name__)

ATTACHMENTS_DIR = pathlib.Path("/app/attachments")

# Maximum extracted text length to avoid token bombs
_MAX_EXTRACT_CHARS = 30000


def _safe_path(filename: str) -> pathlib.Path | None:
    """Resolve attachment path, block traversal."""
    target = (ATTACHMENTS_DIR / filename).resolve()
    try:
        target.relative_to(ATTACHMENTS_DIR.resolve())
    except ValueError:
        log_tool_blocked("attachment_reader", "unknown",
                         f"path traversal attempt: {filename[:100]!r}")
        return None
    if not target.exists():
        return None
    return target


def extract_pdf(path: pathlib.Path) -> str:
    """Extract text from a PDF file."""
    try:
        import pypdf
        reader = pypdf.PdfReader(str(path))
        pages = []
        for i, page in enumerate(reader.pages[:50]):  # cap at 50 pages
            text = page.extract_text() or ""
            if text.strip():
                pages.append(f"--- Page {i+1} ---\n{text}")
        if not pages:
            return "PDF contains no extractable text (might be scanned/image-only)."
        return "\n\n".join(pages)[:_MAX_EXTRACT_CHARS]
    except Exception as exc:
        logger.error(f"PDF extraction failed: {exc}")
        return f"Failed to extract PDF text: {str(exc)[:200]}"


def extract_docx(path: pathlib.Path) -> str:
    """Extract text from a DOCX file."""
    try:
        import docx
        doc = docx.Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        if not paragraphs:
            return "DOCX contains no extractable text."
        return "\n\n".join(paragraphs)[:_MAX_EXTRACT_CHARS]
    except Exception as exc:
        logger.error(f"DOCX extraction failed: {exc}")
        return f"Failed to extract DOCX text: {str(exc)[:200]}"


def extract_xlsx(path: pathlib.Path) -> str:
    """Extract data from an XLSX file as text tables."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        sheets = []
        for name in wb.sheetnames[:10]:  # cap at 10 sheets
            ws = wb[name]
            rows = []
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i > 500:  # cap at 500 rows
                    rows.append("... (truncated)")
                    break
                cells = [str(c) if c is not None else "" for c in row]
                rows.append(" | ".join(cells))
            if rows:
                sheets.append(f"=== Sheet: {name} ===\n" + "\n".join(rows))
        wb.close()
        if not sheets:
            return "XLSX contains no data."
        return "\n\n".join(sheets)[:_MAX_EXTRACT_CHARS]
    except Exception as exc:
        logger.error(f"XLSX extraction failed: {exc}")
        return f"Failed to extract XLSX data: {str(exc)[:200]}"


def extract_image(path: pathlib.Path) -> str:
    """Extract text from an image via GLM-OCR (Ollama), fallback to pytesseract."""
    info = ""
    try:
        from PIL import Image
        img = Image.open(str(path))
        info = f"Image: {img.size[0]}x{img.size[1]} {img.mode} ({img.format})"
    except Exception:
        info = f"Image: {path.name}"

    # Try GLM-OCR via Ollama first (best quality, local, fast)
    try:
        from app.tools.ocr_tool import ocr_from_file
        text = ocr_from_file(str(path))
        if text and not text.startswith("Error") and not text.startswith("OCR failed"):
            return f"{info}\n\nExtracted text (GLM-OCR):\n{text[:_MAX_EXTRACT_CHARS]}"
        logger.debug(f"GLM-OCR returned: {text[:100]}")
    except Exception as e:
        logger.debug(f"GLM-OCR unavailable: {e}")

    # Fallback: pytesseract (traditional OCR)
    try:
        from PIL import Image
        import pytesseract
        img = Image.open(str(path))
        text = pytesseract.image_to_string(img)
        if text.strip():
            return f"{info}\n\nExtracted text (Tesseract OCR):\n{text[:_MAX_EXTRACT_CHARS]}"
        return f"{info}\nOCR found no text in the image."
    except ImportError:
        return f"{info}\nOCR: GLM-OCR not responding and pytesseract not installed."
    except Exception as ocr_exc:
        return f"{info}\nOCR failed: {str(ocr_exc)[:200]}"


# Map MIME types and extensions to extractors
_EXTRACTORS = {
    "application/pdf": extract_pdf,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": extract_docx,
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": extract_xlsx,
    "image/jpeg": extract_image,
    "image/png": extract_image,
    "image/webp": extract_image,
    "image/gif": extract_image,
}

_EXT_MAP = {
    ".pdf": extract_pdf,
    ".docx": extract_docx,
    ".xlsx": extract_xlsx,
    ".xls": extract_xlsx,
    ".jpg": extract_image,
    ".jpeg": extract_image,
    ".png": extract_image,
    ".webp": extract_image,
    ".gif": extract_image,
}


def extract_attachment(filename: str, content_type: str = "") -> str:
    """
    Extract text content from an attachment file.

    Args:
        filename: Relative path within /app/attachments/
        content_type: MIME type (optional, falls back to extension)

    Returns:
        Extracted text content or error message.
    """
    path = _safe_path(filename)
    if path is None:
        return f"Attachment not found: {filename}"

    # Find extractor by MIME type, then by extension
    extractor = _EXTRACTORS.get(content_type)
    if not extractor:
        ext = path.suffix.lower()
        extractor = _EXT_MAP.get(ext)

    if not extractor:
        return (f"Unsupported file type: {content_type or path.suffix}. "
                f"Supported: PDF, DOCX, XLSX, JPG, PNG")

    logger.info(f"Extracting attachment: {filename} ({content_type})")
    return extractor(path)


@tool("read_attachment")
def read_attachment(filename: str, content_type: str = "") -> str:
    """
    Read and extract text from a Signal attachment (PDF, DOCX, XLSX, JPG, PNG).
    filename: the attachment filename as provided in the message
    content_type: optional MIME type hint
    Returns extracted text content.
    """
    return extract_attachment(filename, content_type)
